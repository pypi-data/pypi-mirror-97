import datetime
from typing import List, Optional

import docx
import docx.dml.color
import docx.document
import docx.enum.dml
import docx.enum.section
import docx.enum.style
import docx.enum.table
import docx.enum.text
import docx.oxml.ns
import docx.shared
import docx.styles.styles
import docx.text
import docx.text.paragraph
import docx.section

import vmi


def docx_document(title: str, file: str = None) -> docx.document.Document:
    doccument = docx.Document(file)
    doccument.core_properties.created = datetime.datetime.utcnow()
    doccument.core_properties.modified = datetime.datetime.utcnow()
    doccument.core_properties.title = title
    doccument.core_properties.author = vmi.app.applicationName()
    doccument.core_properties.last_modified_by = vmi.app.applicationName()
    return doccument


def docx_style(document: docx.document.Document, name: str, base: str, font: str = '等线', bold: bool = False,
               size: int = 12, rgb: List[int] = (0, 0, 0), alignment: str = 'LEFT', first_line_indent: float = 0,
               line_spacing: float = 1, space_before: float = 0.5, space_after: float = 0.5, left_indent: float = 0):
    base = document.styles[base]
    style = document.styles.add_style(name, base.type)
    style.base_style = base
    style.quick_style = True

    style.font.name = font
    style.font.bold = bold
    style.font.size = docx.shared.Pt(size)
    style.font.color.rgb = docx.shared.RGBColor(*rgb)
    style.font.no_proof = True
    style.element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), style.font.name)
    alignment = getattr(docx.enum.text.WD_PARAGRAPH_ALIGNMENT, alignment)
    style.paragraph_format.alignment = alignment
    style.paragraph_format.first_line_indent = first_line_indent * style.font.size
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_before = round(space_before * style.font.size)
    style.paragraph_format.space_after = round(space_after * style.font.size)
    style.paragraph_format.left_indent = round(left_indent * style.font.size)

    return style


def docx_section_landscape_A4(section: docx.section.Section):
    section.page_width = docx.shared.Length(10692130)
    section.page_height = docx.shared.Length(7560310)
    section.header_distance = docx.shared.Length(914400)
    section.footer_distance = docx.shared.Length(914400)
    section.orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE


def docx_section_margin_middle(section: docx.section.Section):
    section.top_margin = docx.shared.Length(914400)
    section.bottom_margin = docx.shared.Length(914400)
    section.left_margin = docx.shared.Length(685800)
    section.right_margin = docx.shared.Length(685800)


def docx_section_header(section: docx.section.Section):
    section.header.paragraphs[0].paragraph_format.tab_stops.clear_all()
    section.header.paragraphs[0].paragraph_format.tab_stops.add_tab_stop(
        round(0.5 * (section.page_width - section.left_margin - section.right_margin)),
        docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER)
    section.header.paragraphs[0].paragraph_format.tab_stops.add_tab_stop(
        section.page_width - section.left_margin - section.right_margin,
        docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT)


if __name__ == '__main__':
    d = docx_document('设计确认单')
    docx_section_landscape_A4(d.sections[0])
    docx_section_margin_middle(d.sections[0])
    docx_section_header(d.sections[0])

    yw_title = docx_style(d, '影为标题', 'Normal', size=26, alignment='CENTER')
    yw_header = docx_style(d, '影为页眉', 'Header', size=10)
    yw_body_left = docx_style(d, '影为正文', 'Body Text')
    yw_body_left_indent = docx_style(d, '影为正文首行缩进', 'Body Text', first_line_indent=2)
    yw_body_center = docx_style(d, '影为正文居中', 'Body Text', alignment='CENTER')
    yw_table = docx_style(d, '影为表格', 'Table Grid', alignment='CENTER')

    d.sections[0].header.paragraphs[0].text = '{}\t\t影为医疗科技（上海）有限公司\t\t[表单号]'.format(
        datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S"))
    d.sections[0].header.paragraphs[0].style = yw_header


    def add_row_text(table, texts, styles):
        r = len(table.add_row().table.rows) - 1
        for c in range(len(table.columns)):
            table.cell(r, c).paragraphs[0].text = texts[c]
            table.cell(r, c).paragraphs[0].style = styles[c]


    def add_row_picture(table, file):
        r = len(table.add_row().table.rows) - 1
        table.cell(r, 1).merge(table.cell(r, 3))
        table.cell(r, 1).paragraphs[0].add_run().add_picture(file, width=0.5 * table.cell(r, 1).width)


    top_level_table = d.add_table(1, 4, yw_table)
    top_level_table.autofit = True

    top_level_table.cell(0, 0).merge(top_level_table.cell(0, 3))
    top_level_table.cell(0, 0).paragraphs[0].text = '设计确认单\nDesign Confirmation'
    top_level_table.cell(0, 0).paragraphs[0].style = yw_title

    add_row_text(top_level_table, ['产品名称\nProduct name', '全髋关节置换手术导板\nTHA surgical guide',
                              '序列号\nSerial number', 'THK6KSDJFHKJ'], [yw_body_center] * 4)
    add_row_text(top_level_table, ['型号\nModel', 'TH', '规格\nSpecification', '定制式'], [yw_body_center] * 4)

    add_row_picture(top_level_table, 'c:/py/1.jpg')

    d.save('c:/py/demo.docx')
